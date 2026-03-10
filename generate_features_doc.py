"""Generate FEATURES.docx for Amogha Cafe & Restaurant."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1E, 0x14, 0x0E)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Amogha Cafe & Restaurant', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Complete Feature Documentation')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = desc.add_run('Multi-surface restaurant platform \u2014 online ordering, POS, kiosk, kitchen display, delivery, order tracking, and admin dashboard built on Firebase.')
run2.font.size = Pt(10)
run2.italic = True

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Customer Website (Online Ordering)',
    '2. Admin Dashboard',
    '3. POS Terminal',
    '4. Self-Service Kiosk',
    '5. Kitchen Display System (KDS)',
    '6. Delivery Partner App',
    '7. Order Display Board',
    '8. Order Tracking',
    '9. QR Code Table Ordering (Dine-In)',
    '10. Loyalty Points Lookup',
    '11. REST API (Cloud Functions)',
    '12. Cross-Cutting Concerns',
    '13. Technology Stack',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ============================================================
# 1. CUSTOMER WEBSITE
# ============================================================
doc.add_heading('1. Customer Website (Online Ordering)', level=1)
doc.add_paragraph('Entry point: index.html  |  Modules: src/modules/')

doc.add_heading('1.1 Hero & Navigation', level=2)
bullets([
    'Animated hero slideshow with multiple background images (food photos, chef portraits, restaurant shots)',
    'Header slideshow ticker rotating through today\u2019s specials, bestsellers, Telugu/English proverbs, and promotional messages',
    'Scroll progress bar at the top of the page',
    'Cursor glow trail and floating gold ambient particles for a premium visual feel',
    'Grain texture overlay for a cinematic aesthetic',
    'Preloader with animated gold progress bar and logo',
    'Dark mode toggle (moon icon)',
    'Mobile hamburger menu for responsive navigation',
    'SEO-optimized with Open Graph, Twitter Card, and Schema.org (Restaurant) structured data',
])

doc.add_heading('1.2 Sections', level=2)
bullets([
    'About / Our Story \u2014 brand story, mission, USP list with scroll-reveal animations',
    'Stats counter \u2014 animated counters (10,000+ customers, 50+ dishes, 20+ years, 4.8-star rating)',
    'Meet Our Chefs \u2014 slideshow of chefs (Polina, Usha, Bhavana) with bios and specialties',
    'Today\u2019s Specials \u2014 featured dishes with prices and "Order Now" buttons',
    'Daily Special \u2014 countdown timer section with "Today Only" badge and expiry countdown',
    'Gallery \u2014 image slideshow with dots navigation, lightbox with keyboard controls',
    'Video lightbox \u2014 embedded video player for promotional content',
    'Reviews carousel \u2014 auto-advancing customer reviews with mouse hover pause',
    'Contact section with restaurant info',
])

doc.add_heading('1.3 Menu & Ordering', level=2)
bullets([
    'Live menu sync from Firestore in real-time',
    'Category carousel (Swiggy-style) \u2014 horizontal scrollable category pills with arrow navigation',
    'Menu search with type-ahead filtering',
    'Veg/Non-veg filter indicators on each item',
    'Spice level selector (Mild / Medium / Hot) per item',
    'Add-ons & extras support per menu item',
    'Bestseller and badge tags on menu items',
    'Dynamic pricing rules and happy hour pricing',
    'Delivery banner \u2014 estimated delivery time (30-45 mins) and free delivery threshold (\u20B9500)',
    'Combo Meal Builder \u2014 mix & match starters + main course + bread + drink with 15% discount',
    'Allergen filter ("Safe for Me") \u2014 filter menu based on dietary preferences and allergen alerts',
])

doc.add_heading('1.4 AI-Powered Features', level=2)
bullets([
    'AI Chatbot \u2014 floating chat FAB with AI assistant for menu queries, recommendations, and ordering help',
    'AI Picks For You \u2014 personalized AI-generated food recommendations based on order history',
    'AI Meal Planner \u2014 generate custom meal plans based on preferences',
    'Item pairings \u2014 suggested pairing items (e.g., biryani + raita)',
])

doc.add_heading('1.5 Cart & Checkout', level=2)
bullets([
    'Shopping cart with floating cart icon and item count badge',
    'Floating cart FAB on mobile',
    'Cart modal with item list, quantity controls (+/-), add-ons display, and total',
    'Coupon code system \u2014 apply discount codes at checkout',
    'Gift card support \u2014 redeem gift cards during checkout',
    'Loyalty points redemption \u2014 100 pts = \u20B9100 off, min \u20B9250 order',
    'Payment methods: Razorpay (online), Cash on Delivery, UPI payment links',
    'Delivery fee calculation \u2014 \u20B949 fee, free above \u20B9500',
    'Order notes / special instructions field',
    'Address input for delivery orders',
    'XSS protection \u2014 HTML escaping on all user inputs',
])

doc.add_heading('1.6 Post-Order Features', level=2)
bullets([
    'Order confirmation with order ID and tracking link',
    'WhatsApp bill sharing',
    'Split Bill \u2014 divide order total among 2-10 people with UPI payment link generation per person',
    'Order Again / Reorder \u2014 quick reorder from past orders',
    'Review & Rating prompt \u2014 automatic prompt 1 min after order; rate items 1-5 stars; earns 25 loyalty points',
])

doc.add_heading('1.7 Social & Group Features', level=2)
bullets([
    'Group Ordering \u2014 host creates shared cart, shares link, participants add items, host checks out',
    'Table Reservation \u2014 book tables with date picker, time slots, party size (1-10+), special requests; 7-day advance booking',
])

doc.add_heading('1.8 Subscription Meal Plans', level=2)
bullets([
    'Weekly meal plan subscriptions at discounted prices',
    'Plan catalog loaded from Firestore',
])

doc.add_heading('1.9 User Account', level=2)
bullets([
    'Sign In / Sign Up modal with phone number authentication',
    'Customer profile with dietary preferences, allergen alerts, and saved delivery addresses',
    'Loyalty widget in navbar showing points balance',
    'Badge gallery \u2014 gamification badges displayed in navigation',
])

doc.add_heading('1.10 Gamification & Badges', level=2)
add_table(['Badge', 'Criteria', 'Icon'], [
    ['First Bite', 'Place your first order', '\U0001f37d\ufe0f'],
    ['Regular', '5 orders completed', '\u2b50'],
    ['Foodie', '10 orders completed', '\U0001f3c5'],
    ['Super Fan', '25 orders completed', '\U0001f3c6'],
    ['Explorer', 'Order from all menu categories', '\U0001f5fa\ufe0f'],
    ['Streak Master', '3 consecutive days with orders', '\U0001f525'],
    ['Big Spender', 'Single order over \u20b91000', '\U0001f48e'],
    ['Critic', 'Write 5 reviews', '\U0001f4dd'],
    ['Night Owl', 'Order after 9 PM', '\U0001f989'],
    ['Early Bird', 'Order before 10 AM', '\U0001f426'],
])

doc.add_heading('1.11 Notifications', level=2)
bullets([
    'Browser push notifications \u2014 custom opt-in banner',
    'Order status notifications via browser Notification API',
])

doc.add_heading('1.12 PWA & Offline', level=2)
bullets([
    'PWA install prompt with "Install App" button',
    'Service Worker with offline caching',
    'App manifest for installability',
    'Apple mobile web app support',
])

doc.add_heading('1.13 Ambient Music', level=2)
bullets(['Background music player with two ambient tracks'])

doc.add_heading('1.14 Internationalization (i18n)', level=2)
bullets([
    'Multi-language support via data-i18n attributes',
    'Telugu proverbs and taglines in the header slideshow',
    'Translation constants for UI strings',
])

doc.add_page_break()

# ============================================================
# 2. ADMIN DASHBOARD
# ============================================================
doc.add_heading('2. Admin Dashboard', level=1)
doc.add_paragraph('Entry point: admin/index.html')

doc.add_heading('2.1 Authentication & Access', level=2)
bullets([
    'Admin login with PIN/password protection',
    'Multi-shop switching (dropdown selector for managing multiple locations)',
])

doc.add_heading('2.2 Dashboard Overview', level=2)
bullets([
    'Live order statistics (total orders, today\u2019s count, revenue)',
    'Real-time monitoring',
])

doc.add_heading('2.3 Management Tabs', level=2)
add_table(['Tab', 'Features'], [
    ['Orders', 'Filter by status (all, pending, confirmed, preparing, delivered, cancelled, POS); CSV export'],
    ['Menu', 'Bulk edit, enable/disable items, add new items, manage add-ons and extras'],
    ['Specials', 'Today\u2019s special offers management'],
    ['Daily Special', 'Daily featured item with countdown timer'],
    ['Slideshow', 'Hero slide management (images/videos), customer testimonials'],
    ['Analytics', 'Performance metrics, insights, and reporting'],
    ['Coupons', 'Create and manage discount codes'],
    ['Gift Cards', 'Digital gift card system management'],
    ['Reviews', 'Customer review moderation and management'],
    ['CRM (Customers)', 'Customer database with profiles, order history, loyalty data'],
    ['Expenses', 'Track operational costs; AI-powered OCR bill parsing via Gemini'],
    ['Staff', 'Staff management (roles, schedules)'],
    ['Calendar', 'Schedule and event management'],
    ['Marketing', 'Marketing campaigns and promotions'],
    ['Pricing', 'Dynamic pricing rules management'],
    ['Shops', 'Multi-shop config (name, logo, tagline, theme, categories, admin PIN)'],
    ['Kiosks', 'Kiosk terminal management and monitoring'],
])

doc.add_heading('2.4 Theming', level=2)
bullets([
    'Seasonal themes (e.g., gold/dark holiday theme)',
    'Light/dark mode toggle',
    'Theme selector with decorative noise background',
])

doc.add_page_break()

# ============================================================
# 3. POS TERMINAL
# ============================================================
doc.add_heading('3. POS Terminal', level=1)
doc.add_paragraph('Entry point: pos/index.html')

doc.add_heading('3.1 Authentication', level=2)
bullets([
    'Username/password login for POS operators',
    'Server-side credential validation via Cloud Functions',
])

doc.add_heading('3.2 Core POS Features', level=2)
bullets([
    'Real-time clock display',
    'Sales counter \u2014 real-time daily sales tracking',
    'Category navigation \u2014 quick-access menu category buttons',
    'Smart search \u2014 type-ahead menu search with debouncing',
    'Item frequency tracking \u2014 most-used items highlighted for quick access',
])

doc.add_heading('3.3 Cart & Checkout', level=2)
bullets([
    'Visual cart display with quantity controls (+/-)',
    'Clear cart functionality',
    'Table number assignment',
    'Special notes/instructions (e.g., "extra spicy, no onion")',
    'Multiple payment methods: Cash, UPI, Card',
    'Order summary with real-time total calculation',
    'Bill generation and printing',
])

doc.add_heading('3.4 Customer Integration', level=2)
bullets([
    'Phone number lookup for existing customers',
    'Customer badge display (loyalty tier, points)',
    'Customer name input',
])

doc.add_heading('3.5 Order Management', level=2)
bullets([
    'Recent orders panel with filters (all, POS-only, today, pending)',
    'Session management with auto clock update',
    'Logout functionality',
])

doc.add_page_break()

# ============================================================
# 4. SELF-SERVICE KIOSK
# ============================================================
doc.add_heading('4. Self-Service Kiosk', level=1)
doc.add_paragraph('Entry point: kiosk/index.html')

doc.add_heading('4.1 Multi-Language Support', level=2)
bullets(['English, Telugu (\u0c24\u0c46\u0c32\u0c41\u0c17\u0c41), Hindi (\u0939\u093f\u0902\u0926\u0940)'])

doc.add_heading('4.2 Accessibility', level=2)
bullets([
    'Light/Dark mode toggle',
    'Accessibility mode (adjustable font size with "Aa" button)',
    'Fullscreen functionality',
])

doc.add_heading('4.3 Ordering Features', level=2)
bullets([
    'Voice ordering \u2014 microphone button for voice-based item search',
    'Menu search with autocomplete',
    'Filter chips \u2014 All, Veg, Non-Veg, Popular, Chef\u2019s Pick',
    'Category tabs and sidebar navigation',
    'Item cards with images, names, prices, veg/non-veg indicators, bestseller tags',
    'Quantity controls per item',
])

doc.add_heading('4.4 Cart & Checkout', level=2)
bullets([
    'Floating cart bar (item count + total)',
    'Modal checkout sheet',
    'Customer info: name, phone number, table number',
    'QR code scanner for table identification',
    'Loyalty points display and redemption option',
    '"Order Again" section with recent items',
    'Payment options: Online payment or Pay at Counter',
])

doc.add_heading('4.5 Order Confirmation', level=2)
bullets([
    'Success screen with confetti animation',
    'Order ID, item count, total, payment method',
    'Wait time estimation',
    'WhatsApp bill sharing option',
])

doc.add_page_break()

# ============================================================
# 5. KITCHEN DISPLAY SYSTEM
# ============================================================
doc.add_heading('5. Kitchen Display System (KDS)', level=1)
doc.add_paragraph('Entry point: kitchen/index.html')

doc.add_heading('5.1 Authentication', level=2)
bullets(['Kitchen PIN-based secure access'])

doc.add_heading('5.2 Dashboard Header', level=2)
bullets([
    'Brand display with LIVE status indicator',
    'Active orders count, completed orders count, average prep time, revenue display',
    'Peak traffic visualization, current time display',
])

doc.add_heading('5.3 View Modes', level=2)
add_table(['View', 'Description'], [
    ['Board View', 'Kanban-style columns: NEW \u2192 COOKING \u2192 READY'],
    ['Expo View', 'Detailed order layout for expeditor station'],
    ['Customer View', 'Customer-facing display showing order status'],
    ['Batch View', 'Batch operations for high-volume cooking'],
    ['Analytics View', 'Performance metrics and kitchen efficiency'],
])

doc.add_heading('5.4 Station Filtering', level=2)
bullets(['All Stations, Starters, Curries, Biryanis, Grill & Kebabs, Noodles & Rice, Rotis & Naan'])

doc.add_heading('5.5 Keyboard Shortcuts', level=2)
add_table(['Key', 'Action'], [
    ['/', 'Search orders'],
    ['S', 'Toggle sound alerts'],
    ['V', 'Toggle voice notifications'],
    ['C', 'Open chat/communication'],
    ['G', 'Staff management'],
    ['I', 'Inventory tracking'],
    ['L', 'Table management'],
    ['R', 'Reports'],
    ['H', 'Order history'],
    ['T', 'Theme toggle'],
    ['F', 'Fullscreen mode'],
    ['?', 'Keyboard shortcuts help'],
])

doc.add_heading('5.6 Advanced Features', level=2)
bullets([
    '86 Management \u2014 mark items as sold out ("86\u2019d") across all surfaces',
    'All-Day Count \u2014 track item counts across the full day',
    'Zoom controls (adjust display size)',
    'Batch operations \u2014 START ALL / DONE ALL buttons',
    'Order flash \u2014 visual and audio alerts for new orders',
    'Screensaver mode \u2014 idle screen showing stats',
    'Celebration confetti on order completion',
    'Network status indicator',
])

doc.add_page_break()

# ============================================================
# 6. DELIVERY PARTNER APP
# ============================================================
doc.add_heading('6. Delivery Partner App', level=1)
doc.add_paragraph('Entry point: delivery/index.html')

doc.add_heading('6.1 Authentication', level=2)
bullets(['Phone number + PIN login with server-side validation'])

doc.add_heading('6.2 Tab Navigation', level=2)
add_table(['Tab', 'Features'], [
    ['Available Orders', 'Unclaimed delivery jobs with order details, "Accept Delivery" button'],
    ['Active Delivery', 'Current delivery status, customer info, address with GPS nav, phone with call button, item breakdown, "Mark as Delivered"'],
    ['Earnings', 'Earnings grid, per-order commission (\u20b949/order)'],
    ['Delivery History', 'Past deliveries log'],
])

doc.add_heading('6.3 Navigation & Tracking', level=2)
bullets([
    'GPS tracking \u2014 start/stop location monitoring',
    'Google Maps integration \u2014 one-tap navigation to delivery address',
    'Direct call \u2014 one-tap customer calling',
    'Live order notification badges',
    'Connectivity status indicator',
])

doc.add_page_break()

# ============================================================
# 7. ORDER DISPLAY BOARD
# ============================================================
doc.add_heading('7. Order Display Board', level=1)
doc.add_paragraph('Entry point: display/index.html')
bullets([
    '"Now Serving" section \u2014 rotating carousel of ready orders with pagination dots',
    'Three-column order board \u2014 Preparing | Waiting | Ready (with count per status)',
    'Order cards with order number and customer name',
    'Statistics bar: completed today, average prep time, active orders',
    'LIVE status indicator, current time and date',
    'Particle background animation',
    'Toast notifications and flash overlay for "Order Ready" events',
    'Fullscreen mode, optimized for large displays/TVs',
    'Real-time Firebase synchronization',
])

# ============================================================
# 8. ORDER TRACKING
# ============================================================
doc.add_heading('8. Order Tracking', level=1)
doc.add_paragraph('Entry point: track/index.html')
bullets([
    'Order lookup by order ID',
    '4-step visual stepper: Placed \u2192 Confirmed \u2192 Cooking \u2192 Delivered',
    'Status card with icon, text, and description',
    'Estimated completion time with countdown timer and progress bar',
    'Live delivery tracking \u2014 Google Maps with driver distance and real-time location',
    'Item-by-item breakdown (quantity, price, total)',
    'Delivery address, notes, payment method badge (COD/Online)',
    'Help & contact: direct phone call and WhatsApp buttons',
    'Confetti celebration on order completion',
    'Error / cancelled / loading states handled',
])

# ============================================================
# 9. QR CODE TABLE ORDERING
# ============================================================
doc.add_heading('9. QR Code Table Ordering (Dine-In)', level=1)
doc.add_paragraph('Entry point: qr/index.html')
bullets([
    'Automatic table number parsing from QR code URL parameter',
    'Table badge display',
    'Menu browsing with category tabs, search bar, Veg/Non-Veg filters',
    'Menu items with images, names, prices, veg/non-veg indicators, bestseller badges',
    'Floating cart bar with modal checkout',
    'Customer name, phone (10-digit), and special notes input',
    'Payment: "Pay at Counter" (dine-in model)',
    'Order confirmation with table number, order ID, "Track Order" link, "Place Another Order" button',
])

doc.add_page_break()

# ============================================================
# 10. LOYALTY POINTS LOOKUP
# ============================================================
doc.add_heading('10. Loyalty Points Lookup', level=1)
doc.add_paragraph('Entry point: loyalty/index.html')
bullets([
    'Phone number lookup for loyalty account',
    'Current points balance display',
    'Visual progress bar (0\u2013100 points goal)',
    'Customer statistics: total visits, total spent, last visit date',
])

doc.add_heading('Program Rules', level=2)
add_table(['Rule', 'Detail'], [
    ['Earning rate', '\u20b92000 spent = 100 points'],
    ['Redemption', '100 points = \u20b9100 off'],
    ['Minimum order', '\u20b9250 for redemption'],
    ['Expiry', 'Points never expire'],
    ['Requirement', 'Phone number at checkout'],
])

# ============================================================
# 11. REST API
# ============================================================
doc.add_heading('11. REST API (Cloud Functions)', level=1)
doc.add_paragraph('Entry point: functions/index.js  |  Runtime: Firebase Cloud Functions + Express.js')

doc.add_heading('Public Endpoints', level=2)
add_table(['Method', 'Path', 'Description'], [
    ['GET', '/menu', 'Full available menu with categories, prices, veg/non-veg flags'],
    ['GET', '/specials', 'Today\u2019s specials'],
    ['POST', '/order', 'Place a COD order with server-side price validation'],
    ['GET', '/order/:id', 'Track an order by ID'],
])

doc.add_heading('AI-Powered Endpoints (Gemini 2.0 Flash)', level=2)
add_table(['Method', 'Path', 'Description'], [
    ['POST', '/chat', 'AI chatbot conversation'],
    ['POST', '/smart-search', 'AI-powered menu search'],
    ['POST', '/recommend', 'Personalized recommendations'],
    ['POST', '/meal-plan', 'AI meal plan generation'],
    ['POST', '/smart-combo', 'AI combo suggestions'],
    ['POST', '/summarize-reviews', 'AI review summarization'],
    ['POST', '/parse-bill', 'OCR bill parsing via Gemini Vision (JPG, PNG, WebP, PDF up to 4MB)'],
])

doc.add_heading('Admin Endpoints (Require API Key)', level=2)
add_table(['Method', 'Path', 'Description'], [
    ['POST', '/notify', 'Send notifications'],
    ['POST', '/analytics-query', 'Natural language analytics queries'],
    ['POST', '/forecast', 'Demand forecasting'],
    ['POST', '/menu-insights', 'AI-generated menu insights'],
])

doc.add_heading('Authentication Endpoints', level=2)
add_table(['Method', 'Path', 'Description'], [
    ['POST', '/auth/kiosk-login', 'POS/Kiosk terminal server-side authentication'],
    ['POST', '/auth/delivery-login', 'Delivery partner server-side authentication'],
])

doc.add_heading('Security', level=2)
bullets([
    'Rate limiting \u2014 30 requests/minute per IP on AI endpoints',
    'Admin auth middleware \u2014 API key validation via x-api-key header',
    'CORS restricted to amogha-cafe.web.app, amogha-cafe.firebaseapp.com, amoghahotels.com',
    'Server-side price validation \u2014 order prices verified against menu DB',
    'Input sanitization \u2014 control character stripping and length limits before Gemini prompts',
    'Menu cache \u2014 10-minute in-memory cache to reduce Firestore reads',
])

doc.add_page_break()

# ============================================================
# 12. CROSS-CUTTING CONCERNS
# ============================================================
doc.add_heading('12. Cross-Cutting Concerns', level=1)

doc.add_heading('12.1 Multi-Tenancy', level=2)
bullets([
    'Multiple shop support (Amogha Cafe, Tea Shop)',
    'Shop-specific menus, themes, categories, admin PINs, taglines',
    'Shop switching in admin dashboard',
])

doc.add_heading('12.2 Firebase Services', level=2)
add_table(['Service', 'Usage'], [
    ['Firestore', 'Menu, orders, users, reviews, loyalty, coupons, gift cards, reservations, shops, kiosks, expenses'],
    ['Firebase Hosting', 'Static file hosting with custom domain'],
    ['Cloud Functions', 'REST API, AI integrations, server-side auth'],
    ['Firestore Rules', 'Security rules for data access control'],
    ['Storage Rules', 'File upload access control'],
])

doc.add_heading('12.3 Mobile Apps (Capacitor)', level=2)
bullets([
    'iOS and Android build support',
    'Dedicated build scripts for kiosk and POS Android apps',
    'Asset path fixing for standalone builds',
])

doc.add_heading('12.4 Performance', level=2)
bullets([
    'Font preloading with non-blocking load',
    'DNS prefetch for Firebase, Google Fonts, Razorpay',
    'Critical CSS inlined for instant preloader render',
    'In-memory menu cache in Cloud Functions (10-min TTL)',
    'Debounced search inputs',
])

doc.add_heading('12.5 Security', level=2)
bullets([
    'XSS prevention via HTML escaping helpers',
    'Prompt injection prevention for Gemini',
    'Server-side price validation on orders',
    'Rate limiting on AI endpoints',
    'Admin API key authentication',
    'CORS restrictions',
])

doc.add_heading('12.6 Accessibility', level=2)
bullets([
    'Multi-language (English, Telugu, Hindi)',
    'Font size adjustment mode',
    'Dark/light theme toggle',
    'Keyboard navigation (KDS shortcuts, lightbox arrow keys)',
    'ARIA labels on interactive elements',
])

doc.add_heading('12.7 Offline & PWA', level=2)
bullets([
    'Service workers on all surfaces (website, kiosk, KDS)',
    'Web app manifest for installability',
    'Offline-first caching strategies',
])

doc.add_page_break()

# ============================================================
# 13. TECHNOLOGY STACK
# ============================================================
doc.add_heading('13. Technology Stack', level=1)
add_table(['Layer', 'Technology'], [
    ['Frontend', 'Vanilla JavaScript (ES modules), HTML5, CSS3 (CSS variables)'],
    ['Backend', 'Firebase Cloud Functions, Express.js'],
    ['Database', 'Firebase Firestore'],
    ['AI', 'Google Vertex AI (Gemini 2.0 Flash)'],
    ['Payments', 'Razorpay, UPI, Cash on Delivery'],
    ['Mobile', 'Capacitor (iOS + Android)'],
    ['Maps', 'Google Maps (delivery tracking)'],
    ['Build', 'Vite (dev/build), custom shell scripts (kiosk/POS builds)'],
    ['Testing', 'Vitest (2,509 unit tests), Playwright (E2E)'],
    ['Hosting', 'Firebase Hosting'],
    ['Fonts', 'Google Fonts (Cormorant Garamond, Playfair Display, Poppins, Noto Serif Telugu)'],
])

# ============================================================
# MENU CATALOG
# ============================================================
doc.add_heading('Menu Catalog', level=1)
doc.add_paragraph('383+ items across 15+ categories:')
add_table(['Category', 'Item Count'], [
    ['Biryanis', '46'],
    ['Rice / Fried Rice', '41'],
    ['Noodles', '43'],
    ['Starters', '53'],
    ['Curries', '42'],
    ['Pulao', '46'],
    ['Soups', '20'],
    ['Tiffins', '43'],
    ['Breads, Rolls, Sides, Beverages, Sweets', 'Various'],
    ['Tea Shop items', '10'],
])

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Generated from codebase analysis \u2014 March 2026')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.italic = True

# Save
doc.save('/home/user/amogha-cafe/docs/Amogha_Cafe_Features.docx')
print('Done: docs/Amogha_Cafe_Features.docx')
